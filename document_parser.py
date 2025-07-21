"""
Advanced document parser for Sutra application.
Extracts structured content from documents, including text, tables, images, and specialized content.
"""

import os
import uuid
import base64
import tempfile
import io
import re
from typing import Dict, List, Any, Optional, Tuple, BinaryIO, Union
from pathlib import Path
import json

import pytesseract
from PIL import Image
import pandas as pd
import pdfplumber
import docx
from pdf2image import convert_from_path
import numpy as np
from sklearn.cluster import DBSCAN

# Optional heavy libraries – imported lazily or wrapped in try/except so they’re not hard requirements at runtime
try:
    import fitz  # PyMuPDF
except ImportError:  # pragma: no cover
    fitz = None

try:
    import camelot
except ImportError:  # pragma: no cover
    camelot = None

try:
    import tabula
except ImportError:  # pragma: no cover
    tabula = None

try:
    import easyocr
except ImportError:  # pragma: no cover
    easyocr = None

try:
    import layoutparser as lp
except ImportError:  # pragma: no cover
    lp = None

try:
    import cv2  # OpenCV for preprocessing
except ImportError:  # pragma: no cover
    cv2 = None

try:
    import ocrmypdf
except ImportError:  # pragma: no cover
    ocrmypdf = None

from document_classifier import detect_document_type, get_extraction_hints


class BoundingBox:
    """Represents a rectangular region in a document."""
    
    def __init__(self, x0: float, y0: float, x1: float, y1: float, page: int = 0):
        """
        Initialize bounding box.
        
        Args:
            x0: Left coordinate
            y0: Top coordinate
            x1: Right coordinate
            y1: Bottom coordinate
            page: Page number (0-indexed)
        """
        self.x0 = x0
        self.y0 = y0
        self.x1 = x1
        self.y1 = y1
        self.page = page
        
    def area(self) -> float:
        """Calculate area of the bounding box."""
        return (self.x1 - self.x0) * (self.y1 - self.y0)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary representation."""
        return {
            "x0": self.x0,
            "y0": self.y0,
            "x1": self.x1,
            "y1": self.y1,
            "page": self.page
        }
    
    def overlaps(self, other: 'BoundingBox') -> bool:
        """Check if this box overlaps with another box."""
        return (
            self.page == other.page and
            self.x0 < other.x1 and
            self.x1 > other.x0 and
            self.y0 < other.y1 and
            self.y1 > other.y0
        )
    
    def contains(self, other: 'BoundingBox') -> bool:
        """Check if this box fully contains another box."""
        return (
            self.page == other.page and
            self.x0 <= other.x0 and
            self.y0 <= other.y0 and
            self.x1 >= other.x1 and
            self.y1 >= other.y1
        )
    
    @classmethod
    def merge(cls, boxes: List['BoundingBox']) -> 'BoundingBox':
        """Merge multiple boxes into a single encompassing box."""
        if not boxes:
            return None
        
        page = boxes[0].page
        x0 = min(box.x0 for box in boxes)
        y0 = min(box.y0 for box in boxes)
        x1 = max(box.x1 for box in boxes)
        y1 = max(box.y1 for box in boxes)
        
        return cls(x0, y0, x1, y1, page)


class StructuredElement:
    """Class representing a structured element extracted from a document."""
    
    def __init__(
        self,
        section_id: str,
        content_type: str,
        text: str = "",
        image_snippets: List[str] = None,
        metadata: Dict[str, Any] = None,
        notes: str = "",
        bounding_box: BoundingBox = None,
        page: int = None
    ):
        """
        Initialize a structured element.
        
        Args:
            section_id: Unique identifier for this section
            content_type: Type of content (text, table, image, formula, question, entry)
            text: Extracted text content
            image_snippets: List of image file paths for non-text content
            metadata: Additional metadata about this element
            notes: Processing notes (OCR used, etc)
            bounding_box: Position of this element in the document
            page: Page number where this element appears
        """
        self.section_id = section_id
        self.content_type = content_type
        self.text = text
        self.image_snippets = image_snippets or []
        self.metadata = metadata or {}
        self.notes = notes
        self.bounding_box = bounding_box
        
        # Set page number
        if page is not None:
            self.page = page
        elif bounding_box is not None:
            self.page = bounding_box.page
        else:
            self.page = None
    
    def to_dict(self) -> Dict[str, Any]:
        """
        Convert this element to a dictionary representation.
        
        Returns:
            Dictionary representation of this structured element
        """
        result = {
            "section_id": self.section_id,
            "type": self.content_type,
            "content": self.text,
        }
        
        # Add page if available
        if self.page is not None:
            result["page"] = self.page + 1  # Convert to 1-indexed for output
            
        # Add bounding box if available
        if self.bounding_box:
            result["bbox"] = self.bounding_box.to_dict()
            
        # Add image snippets if available
        if self.image_snippets:
            result["snippets"] = self.image_snippets
            
        # Add type-specific fields from metadata
        if self.metadata:
            for key, value in self.metadata.items():
                if key not in result:
                    result[key] = value
                    
        # Special handling for different content types
        if self.content_type == "question":
            if "options" in self.metadata:
                result["options"] = self.metadata["options"]
            if "answer" in self.metadata:
                result["answer"] = self.metadata["answer"]
            if "solution" in self.metadata:
                result["solution"] = self.metadata["solution"]
                
        elif self.content_type == "table":
            if "table_data" in self.metadata:
                result["table_data"] = self.metadata["table_data"]
                
        # Include processing notes if present
        if self.notes:
            result["notes"] = self.notes
            
        return result


class DocumentParser:
    """Class for parsing documents into structured content."""
    
    def __init__(self, output_dir: str = None):
        """
        Initialize the document parser.
        
        Args:
            output_dir: Directory to store extracted images
        """
        self.output_dir = output_dir or tempfile.mkdtemp()
        os.makedirs(self.output_dir, exist_ok=True)
        
        # Lazy EasyOCR reader (initialised on first use)
        self._easyocr_reader = None

        # Patterns for recognizing question elements
        self.question_patterns = {
            'question_id': r'(?:^|\s)(?:Question|Q|Problem|Exercise)\s*(?:#\s*)?(\d+[\.\):])',
            'multiple_choice_options': r'(?:^|\s)([A-Da-d])[\.:\)]\s+(.*?)(?=$|(?:\s+[A-Da-d][\.:\)])\s+)',
            'answer_marker': r'(?:Answer|Answers|Solution|Solutions):.*',
            'math_block': r'(?:\$\$.*?\$\$|\\\[.*?\\\])',
            'formula': r'(?:\$.*?\$|\\begin\{equation\}.*?\\end\{equation\})',
            'multiple_choice_indicator': r'(?:Select one|Choose one|Select all that apply|Choose all that apply|Mark all correct)'
        }
        
    def parse_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse a PDF file into structured content.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of structured elements
        """
        # Extract text using pdfplumber
        elements = []
        section_counter = 0
        
        # ------------------------------------------------------------------
        # Pre-process scanned PDFs with OCRmyPDF to add text layer
        # This enables Camelot and pdfplumber to work on image-based pages
        # ------------------------------------------------------------------

        preprocessed_path = file_path
        if ocrmypdf is not None:
            try:
                with pdfplumber.open(file_path) as pdf:
                    # Check first page for text; if empty, assume scanned
                    first_page_text = pdf.pages[0].extract_text() or ""
                    if not first_page_text.strip():
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_ocr_file:
                            temp_ocr_path = temp_ocr_file.name
                            temp_ocr_file.close()  # Close so OCRmyPDF can write
                        ocrmypdf.ocr(file_path, temp_ocr_path, deskew=True, clean=True, force_ocr=True)
                        preprocessed_path = temp_ocr_path
            except Exception:
                # If OCRmyPDF fails, proceed with original file
                preprocessed_path = file_path

        try:
            # Convert entire PDF to images for later use if needed
            try:
                pdf_images = convert_from_path(preprocessed_path, dpi=300)
            except Exception:
                pdf_images = []
            page_images = []

            # Save each page as image
            for i, img in enumerate(pdf_images):
                img_path = os.path.join(self.output_dir, f"page_{i+1}.png")
                img.save(img_path)
                page_images.append(img_path)
            
            with pdfplumber.open(preprocessed_path) as pdf:
                # First, detect document type from first few pages
                sample_text = ""
                for i, page in enumerate(pdf.pages[:min(3, len(pdf.pages))]):
                    try:
                        sample_text += page.extract_text() or ""
                    except:
                        pass
                    if i >= 2:  # Limit to first 3 pages
                        break
                
                # Detect document type
                doc_info = detect_document_type(sample_text)
                document_type = doc_info["document_type"]
                extraction_hints = get_extraction_hints(document_type)
                
                # Add document type information
                elements.append(
                    StructuredElement(
                        section_id="document_info",
                        content_type="metadata",
                        text="",
                        metadata={
                            "document_type": document_type,
                            "confidence": doc_info.get("confidence", 0),
                            "extraction_hints": extraction_hints
                        }
                    ).to_dict()
                )
                
                # Process each page
                for i, page in enumerate(pdf.pages):
                    try:
                        # 0. Ensure we have a rasterised image for this page
                        if i >= len(pdf_images):
                            try:
                                page_img_pil = page.to_image(resolution=300).original
                                img_path_fallback = os.path.join(self.output_dir, f"page_{i+1}.png")
                                page_img_pil.save(img_path_fallback)
                                pdf_images.append(page_img_pil)
                                page_images.append(img_path_fallback)
                            except Exception:
                                # Create a blank image placeholder so downstream code has a path
                                blank_img = Image.new("RGB", (int(page.width), int(page.height)), color="white")
                                img_path_fallback = os.path.join(self.output_dir, f"page_{i+1}.png")
                                blank_img.save(img_path_fallback)
                                page_images.append(img_path_fallback)

                        # 1. Extract page text – if empty, fall back to OCR
                        text = page.extract_text() or ""
                        if not text.strip():
                            text = self._ocr_page_image(page_images[i] if i < len(page_images) else None)
                        
                        # 2. Extract tables
                        tables = page.extract_tables()
                        table_bboxes = []
                        
                        # Process tables first
                        for j, table in enumerate(tables):
                            if table:
                                # Convert table to text and DataFrame
                                table_text = self._table_to_text(table)
                                try:
                                    table_df = pd.DataFrame(table[1:], columns=table[0])
                                    table_json = table_df.to_json(orient="records")
                                except:
                                    table_json = json.dumps(table)
                                
                                # Get table bounding box
                                table_bbox = BoundingBox(0, 0, page.width, page.height, i)  # Default full page
                                
                                # Try to get actual table bounds if available
                                try:
                                    tables_with_bbox = page.find_tables()
                                    if j < len(tables_with_bbox):
                                        tb = tables_with_bbox[j]
                                        table_bbox = BoundingBox(tb.bbox[0], tb.bbox[1], tb.bbox[2], tb.bbox[3], i)
                                        table_bboxes.append(table_bbox)
                                except:
                                    pass
                                
                                section_id = f"page_{i+1}_table_{j+1}"
                                elements.append(
                                    StructuredElement(
                                        section_id=section_id,
                                        content_type="table",
                                        text=table_text,
                                        metadata={
                                            "page_number": i+1,
                                            "table_number": j+1,
                                            "table_data": json.loads(table_json)
                                        },
                                        bounding_box=table_bbox,
                                        page=i
                                    ).to_dict()
                                )
                        
                        # 3. Extract questions and other structured content
                        if document_type in ["exam_paper", "multiple_choice"]:
                            question_elements = self._extract_questions(text, page_images[i], i)
                            elements.extend([el.to_dict() for el in question_elements])
                            
                        elif document_type == "academic_paper":
                            # Extract formulas and diagrams
                            formula_elements = self._extract_formulas(text, page_images[i], i)
                            elements.extend([el.to_dict() for el in formula_elements])
                            
                        # 4. Extract text blocks that don't overlap with tables/questions
                        if text:
                            # For now, just add the text as a block
                            # In a more sophisticated version, we would segment by paragraphs
                            section_id = f"page_{i+1}_text"
                            elements.append(
                                StructuredElement(
                                    section_id=section_id,
                                    content_type="text",
                                    text=text,
                                    metadata={"page_number": i+1},
                                    page=i
                                ).to_dict()
                            )
                            
                        # 5. Look for regions that might contain images or formulas
                        # Use page image for this
                        image_elements = self._extract_image_elements(page, page_images[i], i, table_bboxes)
                        elements.extend([el.to_dict() for el in image_elements])
                        
                    except Exception as e:
                        # Add an error element for failed pages
                        section_id = f"page_{i+1}_error"
                        elements.append(
                            StructuredElement(
                                section_id=section_id,
                                content_type="error",
                                text="",
                                notes=f"Failed to process page: {str(e)}",
                                page=i
                            ).to_dict()
                        )
        
            # 6. Additional embedded images using PyMuPDF (captures images not seen by pdfplumber)
            if fitz is not None:
                try:
                    pymupdf_images = self._extract_embedded_images_pymupdf(preprocessed_path)
                    elements.extend([el.to_dict() for el in pymupdf_images])
                except Exception:
                    pass

            # 7. Additional tables via Camelot if pdfplumber found none
            if camelot is not None and not any(el for el in elements if el.get("type") == "table"):
                try:
                    camelot_tables = self._extract_tables_camelot(preprocessed_path)
                    elements.extend([el.to_dict() for el in camelot_tables])
                except Exception:
                    pass

        except Exception as e:
            # Add a document-level error
            elements.append(
                StructuredElement(
                    section_id="document_error",
                    content_type="error",
                    text="",
                    notes=f"Failed to process document: {str(e)}"
                ).to_dict()
            )
        
        # Clean up temp OCR file if created
        if preprocessed_path != file_path:
            try:
                Path(preprocessed_path).unlink(missing_ok=True)
            except Exception:
                pass

        return elements
    
    def _extract_questions(self, text: str, page_image_path: str, page_num: int) -> List[StructuredElement]:
        """
        Extract questions from text.
        
        Args:
            text: Text to analyze
            page_image_path: Path to the page image
            page_num: Page number
            
        Returns:
            List of structured elements representing questions
        """
        elements = []
        
        # Look for question patterns
        question_matches = list(re.finditer(self.question_patterns['question_id'], text, re.MULTILINE))
        
        # If no questions found, return empty list
        if not question_matches:
            return elements
            
        # Extract question blocks
        for i, match in enumerate(question_matches):
            # Determine question boundaries
            start_pos = match.start()
            end_pos = len(text)
            
            # If there's another question after this one, use that as boundary
            if i < len(question_matches) - 1:
                end_pos = question_matches[i + 1].start()
                
            # Extract question text
            question_text = text[start_pos:end_pos].strip()
            question_id = match.group(1).strip()
            
            # Look for multiple choice options
            options = {}
            option_matches = list(re.finditer(self.question_patterns['multiple_choice_options'], question_text, re.MULTILINE))
            
            for opt_match in option_matches:
                letter = opt_match.group(1)
                option_text = opt_match.group(2).strip()
                options[letter] = option_text
                
            # Look for answer (if present)
            answer = None
            answer_match = re.search(self.question_patterns['answer_marker'], question_text)
            if answer_match:
                answer_text = question_text[answer_match.start():].strip()
                # Try to identify the answer from the text
                for letter in options.keys():
                    if re.search(rf'\b{letter}\b', answer_text):
                        if answer is None:
                            answer = []
                        answer.append(letter)
                        
            # Create the question element
            section_id = f"page_{page_num+1}_question_{question_id}"
            
            # Check if there are math formulas in the question
            contains_formula = bool(re.search(self.question_patterns['formula'], question_text))
            
            if contains_formula or not options:
                # For questions with formulas or no detected options, include the image snippet
                # Crop the page image to get just this question (if possible)
                snippet_path = f"question_{page_num+1}_{question_id.replace('.', '_')}.png"
                full_snippet_path = os.path.join(self.output_dir, snippet_path)
                
                try:
                    # Here we would ideally crop the image to the question bounds
                    # For simplicity, we'll just use the full page image for now
                    img = Image.open(page_image_path)
                    img.save(full_snippet_path)
                    
                    elements.append(
                        StructuredElement(
                            section_id=section_id,
                            content_type="question",
                            text=question_text,
                            image_snippets=[snippet_path],
                            metadata={
                                "options": options,
                                "answer": answer,
                                "question_id": question_id
                            },
                            page=page_num
                        )
                    )
                except Exception:
                    # Fall back to text only
                    elements.append(
                        StructuredElement(
                            section_id=section_id,
                            content_type="question",
                            text=question_text,
                            metadata={
                                "options": options,
                                "answer": answer,
                                "question_id": question_id
                            },
                            page=page_num
                        )
                    )
            else:
                # Regular question with text options
                elements.append(
                    StructuredElement(
                        section_id=section_id,
                        content_type="question",
                        text=question_text,
                        metadata={
                            "options": options,
                            "answer": answer,
                            "question_id": question_id
                        },
                        page=page_num
                    )
                )
                
        return elements
    
    def _extract_formulas(self, text: str, page_image_path: str, page_num: int) -> List[StructuredElement]:
        """
        Extract mathematical formulas from text.
        
        Args:
            text: Text to analyze
            page_image_path: Path to the page image
            page_num: Page number
            
        Returns:
            List of structured elements representing formulas
        """
        elements = []
        
        # Look for formula patterns (both inline and block)
        formula_matches = list(re.finditer(self.question_patterns['formula'], text, re.MULTILINE))
        math_block_matches = list(re.finditer(self.question_patterns['math_block'], text, re.MULTILINE))
        
        all_matches = formula_matches + math_block_matches
        
        for i, match in enumerate(all_matches):
            formula_text = match.group(0)
            
            # Create a snippet for each formula
            snippet_path = f"formula_{page_num+1}_{i+1}.png"
            full_snippet_path = os.path.join(self.output_dir, snippet_path)
            
            try:
                # Here we would ideally crop the image to the formula bounds
                # For simplicity, we'll just use the full page image for now
                img = Image.open(page_image_path)
                img.save(full_snippet_path)
                
                section_id = f"page_{page_num+1}_formula_{i+1}"
                elements.append(
                    StructuredElement(
                        section_id=section_id,
                        content_type="math_block" if match in math_block_matches else "formula",
                        text=formula_text,
                        image_snippets=[snippet_path],
                        page=page_num
                    )
                )
            except Exception:
                # Skip if we can't create the snippet
                pass
                
        return elements
    
    def _extract_image_elements(
        self,
        page: 'pdfplumber.page.Page',
        page_image_path: str,
        page_num: int,
        exclude_bboxes: List[BoundingBox]
    ) -> List[StructuredElement]:
        """
        Extract image elements from a page that don't overlap with tables.
        
        Args:
            page: pdfplumber Page object for this page
            page_image_path: Path to the rasterised full-page image (for fallback use)
            page_num: Page number
            exclude_bboxes: Bounding boxes to exclude
            
        Returns:
            List of structured elements
        """
        elements: List[StructuredElement] = []

        # 1. Attempt to use pdfplumber's detected images on the page
        try:
            images_on_page = getattr(page, "images", [])

            for idx, img_obj in enumerate(images_on_page):
                # Build bounding box for this image
                bbox = BoundingBox(
                    img_obj.get("x0", 0),
                    img_obj.get("top", img_obj.get("y0", 0)),
                    img_obj.get("x1", page.width),
                    img_obj.get("bottom", img_obj.get("y1", page.height)),
                    page_num,
                )

                # Skip if overlaps with any excluded bbox (e.g., tables)
                if any(bbox.overlaps(ex_bb) for ex_bb in exclude_bboxes):
                    continue

                # Crop the image region from the PDF and save
                snippet_path = f"page_{page_num+1}_image_{idx+1}.png"
                full_snippet_path = os.path.join(self.output_dir, snippet_path)

                try:
                    cropped_page = page.crop((bbox.x0, bbox.y0, bbox.x1, bbox.y1))
                    cropped_page.to_image(resolution=300).save(full_snippet_path, format="PNG")
                except Exception:
                    # Fallback: crop the rasterised page image if PDF crop fails
                    try:
                        page_img = Image.open(page_image_path)
                        scale_x = page_img.width / float(page.width)
                        scale_y = page_img.height / float(page.height)

                        crop_box_px = (
                            int(bbox.x0 * scale_x),
                            int(bbox.y0 * scale_y),
                            int(bbox.x1 * scale_x),
                            int(bbox.y1 * scale_y),
                        )
                        page_img.crop(crop_box_px).save(full_snippet_path)
                    except Exception:
                        # If even the fallback fails, skip this image
                        continue

                section_id = f"page_{page_num+1}_image_{idx+1}"
                elements.append(
                    StructuredElement(
                        section_id=section_id,
                        content_type="image",
                        text="",
                        image_snippets=[snippet_path],
                        bounding_box=bbox,
                        page=page_num,
                        notes="Extracted image/diagram block",
                    )
                )

        except Exception:
            # If pdfplumber detection fails, fall back to full-page snippet below
            images_on_page = []

        # 2. Fallback: if no individual images detected, keep the previous behaviour (full page reference)
        if not elements:
            try:
                img = Image.open(page_image_path)
                section_id = f"page_{page_num+1}_content"
                snippet_path = f"content_{page_num+1}.png"
                full_snippet_path = os.path.join(self.output_dir, snippet_path)
                img.save(full_snippet_path)

                elements.append(
                    StructuredElement(
                        section_id=section_id,
                        content_type="image",
                        text="",
                        image_snippets=[snippet_path],
                        notes="Full page image for reference",
                        page=page_num,
                    )
                )
            except Exception:
                pass

        return elements
    
    def parse_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse a DOCX file into structured content.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of structured elements
        """
        elements = []
        
        try:
            doc = docx.Document(file_path)
            
            # Extract full text first for document classification
            full_text = "\n".join([p.text for p in doc.paragraphs])
            
            # Detect document type
            doc_info = detect_document_type(full_text)
            document_type = doc_info["document_type"]
            extraction_hints = get_extraction_hints(document_type)
            
            # Add document type information
            elements.append(
                StructuredElement(
                    section_id="document_info",
                    content_type="metadata",
                    text="",
                    metadata={
                        "document_type": document_type,
                        "confidence": doc_info.get("confidence", 0),
                        "extraction_hints": extraction_hints
                    }
                ).to_dict()
            )
            
            # Process paragraphs
            current_section = ""
            section_paragraphs = []
            section_count = 0
            
            for i, paragraph in enumerate(doc.paragraphs):
                # Check if this is a heading/section title
                if paragraph.style.name.startswith('Heading'):
                    # If we have collected text from previous section, add it
                    if section_paragraphs:
                        section_text = "\n".join(section_paragraphs)
                        if section_text.strip():
                            section_id = f"section_{section_count}"
                            elements.append(
                                StructuredElement(
                                    section_id=section_id,
                                    content_type="text",
                                    text=section_text,
                                    metadata={"section_title": current_section}
                                ).to_dict()
                            )
                        section_paragraphs = []
                        section_count += 1
                    
                    # Update current section title
                    current_section = paragraph.text
                    section_id = f"heading_{section_count}"
                    elements.append(
                        StructuredElement(
                            section_id=section_id,
                            content_type="heading",
                            text=paragraph.text,
                            metadata={"level": int(paragraph.style.name.replace('Heading', ''))}
                        ).to_dict()
                    )
                else:
                    # Check for question patterns
                    if document_type in ["exam_paper", "multiple_choice"]:
                        question_match = re.search(self.question_patterns['question_id'], paragraph.text)
                        if question_match:
                            # If we have collected text from previous section, add it
                            if section_paragraphs:
                                section_text = "\n".join(section_paragraphs)
                                if section_text.strip():
                                    section_id = f"section_{section_count}"
                                    elements.append(
                                        StructuredElement(
                                            section_id=section_id,
                                            content_type="text",
                                            text=section_text,
                                            metadata={"section_title": current_section}
                                        ).to_dict()
                                    )
                                section_paragraphs = []
                                section_count += 1
                            
                            # Start a new question
                            question_id = question_match.group(1).strip()
                            section_id = f"question_{question_id}"
                            elements.append(
                                StructuredElement(
                                    section_id=section_id,
                                    content_type="question",
                                    text=paragraph.text,
                                    metadata={"question_id": question_id}
                                ).to_dict()
                            )
                            continue
                        
                        # Check for multiple choice options
                        option_match = re.search(r'^([A-Da-d])[\.:\)]\s+(.*)', paragraph.text)
                        if option_match:
                            # Find the last question element
                            last_question = None
                            for el in reversed(elements):
                                if el.get("type") == "question":
                                    last_question = el
                                    break
                            
                            if last_question:
                                # Add this option to the question
                                letter = option_match.group(1)
                                option_text = option_match.group(2).strip()
                                
                                if "options" not in last_question:
                                    last_question["options"] = {}
                                last_question["options"][letter] = option_text
                                continue
                    
                    # Add to current section's paragraphs
                    if paragraph.text.strip():
                        section_paragraphs.append(paragraph.text)
            
            # Add the last section if not empty
            if section_paragraphs:
                section_text = "\n".join(section_paragraphs)
                if section_text.strip():
                    section_id = f"section_{section_count}"
                    elements.append(
                        StructuredElement(
                            section_id=section_id,
                            content_type="text",
                            text=section_text,
                            metadata={"section_title": current_section}
                        ).to_dict()
                    )
            
            # Process tables
            for i, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    rows.append(cells)
                
                # Convert to text and JSON
                table_text = self._table_to_text(rows)
                try:
                    table_df = pd.DataFrame(rows[1:], columns=rows[0])
                    table_json = table_df.to_json(orient="records")
                except:
                    table_json = json.dumps(rows)
                
                section_id = f"table_{i+1}"
                elements.append(
                    StructuredElement(
                        section_id=section_id,
                        content_type="table",
                        text=table_text,
                        metadata={"table_data": json.loads(table_json)}
                    ).to_dict()
                )
            
        except Exception as e:
            # Add a document-level error
            elements.append(
                StructuredElement(
                    section_id="document_error",
                    content_type="error",
                    text="",
                    notes=f"Failed to process document: {str(e)}"
                ).to_dict()
            )
        
        return elements
    
    def parse_image(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse an image file into structured content using OCR.
        
        Args:
            file_path: Path to the image file
            
        Returns:
            List of structured elements
        """
        elements = []
        
        try:
            # Open image
            img = Image.open(file_path)
            
            # Save a copy in output dir
            img_filename = os.path.basename(file_path)
            img_output_path = os.path.join(self.output_dir, img_filename)
            img.save(img_output_path)
            
            # ---------------------------
            # 1. OCR via EasyOCR (preferred)
            # ---------------------------
            ocr_text = ""
            if easyocr is not None:
                try:
                    if self._easyocr_reader is None:
                        self._easyocr_reader = easyocr.Reader(["en"], gpu=False)
                    ocr_text = "\n".join(self._easyocr_reader.readtext(img, detail=0))
                except Exception:
                    ocr_text = ""

            # ---------------------------
            # 2. OCR via pytesseract with preprocessing
            # ---------------------------
            if not ocr_text.strip():
                # Preprocess image for pytesseract if OpenCV is available
                processed_img = img
                if cv2 is not None:
                    try:
                        processed_img = self._preprocess_image_for_ocr(img)
                    except Exception:
                        processed_img = img

                try:
                    ocr_text = pytesseract.image_to_string(processed_img, config="--psm 6")
                except Exception:
                    ocr_text = ""

            if ocr_text and ocr_text.strip():
                # Detect document type
                doc_info = detect_document_type(ocr_text)
                document_type = doc_info["document_type"]
                extraction_hints = get_extraction_hints(document_type)
                
                # Add document type information
                elements.append(
                    StructuredElement(
                        section_id="document_info",
                        content_type="metadata",
                        text="",
                        metadata={
                            "document_type": document_type,
                            "confidence": doc_info.get("confidence", 0),
                            "extraction_hints": extraction_hints
                        }
                    ).to_dict()
                )
                
                # Extract specialized content based on document type
                if document_type in ["exam_paper", "multiple_choice"]:
                    question_elements = self._extract_questions(ocr_text, img_output_path, 0)
                    for el in question_elements:
                        elements.append(el.to_dict())
                
                # Add OCR text
                elements.append(
                    StructuredElement(
                        section_id="image_text",
                        content_type="text",
                        text=ocr_text,
                        notes="Extracted via OCR"
                    ).to_dict()
                )
            else:
                # Add image reference if OCR didn't yield usable text
                elements.append(
                    StructuredElement(
                        section_id="document_info",
                        content_type="metadata",
                        text="",
                        metadata={"document_type": "image"}
                    ).to_dict()
                )
                
                elements.append(
                    StructuredElement(
                        section_id="image_content",
                        content_type="image",
                        text="",
                        image_snippets=[img_filename],
                        notes="OCR did not yield usable text"
                    ).to_dict()
                )
        
        except Exception as e:
            # Add a document-level error
            elements.append(
                StructuredElement(
                    section_id="document_error",
                    content_type="error",
                    text="",
                    notes=f"Failed to process image: {str(e)}"
                ).to_dict()
            )
        
        return elements
    
    def parse_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse any supported file type into structured content.
        
        Args:
            file_path: Path to the file
            
        Returns:
            List of structured elements
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == ".pdf":
            return self.parse_pdf(file_path)
        elif file_ext == ".docx":
            return self.parse_docx(file_path)
        elif file_ext in [".jpg", ".jpeg", ".png"]:
            return self.parse_image(file_path)
        else:
            # Unsupported file type
            return [
                StructuredElement(
                    section_id="file_error",
                    content_type="error",
                    text="",
                    notes=f"Unsupported file type: {file_ext}"
                ).to_dict()
            ]
    
    def _table_to_text(self, table: List[List[str]]) -> str:
        """
        Convert a table to a text representation.
        
        Args:
            table: List of rows, where each row is a list of cell values
            
        Returns:
            String representation of the table
        """
        if not table:
            return ""
        
        # Calculate column widths
        col_widths = [0] * len(table[0])
        for row in table:
            for i, cell in enumerate(row):
                if i < len(col_widths):
                    col_widths[i] = max(col_widths[i], len(str(cell)))
        
        # Format table
        result = []
        for row in table:
            row_text = " | ".join(str(cell).ljust(col_widths[i]) for i, cell in enumerate(row) if i < len(col_widths))
            result.append(row_text)
            
            # Add separator after header
            if row == table[0]:
                result.append("-" * len(row_text))
        
        return "\n".join(result)

    # ------------------------------------------------------------------
    # Helper functions leveraging additional libraries
    # ------------------------------------------------------------------

    def _ocr_page_image(self, page_image_path: Optional[str]) -> str:
        """Run OCR on a page image using EasyOCR (if available) or pytesseract."""
        text = ""
        if not page_image_path:
            return ""

        # Preferred: EasyOCR for better diagram handling
        if easyocr is not None:
            if self._easyocr_reader is None:
                # Initialise once; GPU set to False for wider compatibility
                self._easyocr_reader = easyocr.Reader(["en"], gpu=False)
            try:
                result = self._easyocr_reader.readtext(page_image_path, detail=0)
                text = "\n".join(result)
            except Exception:
                text = ""

        # Fallback to pytesseract
        if not text.strip():
            try:
                img = Image.open(page_image_path)
                text = pytesseract.image_to_string(img)
            except Exception:
                text = ""

        return text

    def _extract_embedded_images_pymupdf(self, file_path: str) -> List[StructuredElement]:
        """Extract images embedded in the PDF using PyMuPDF (fitz)."""
        if fitz is None:
            return []

        elements: List[StructuredElement] = []
        try:
            doc = fitz.open(file_path)
            for page_idx in range(len(doc)):
                page = doc[page_idx]
                img_list = page.get_images(full=True)
                for img_index, img_info in enumerate(img_list):
                    xref = img_info[0]
                    try:
                        pix = fitz.Pixmap(doc, xref)
                        if pix.n > 4:
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                        img_name = f"page_{page_idx+1}_embed_{img_index+1}.png"
                        pix.save(os.path.join(self.output_dir, img_name))
                    except Exception:
                        continue

                    # Try to get bounding box
                    try:
                        bbox_rect = page.get_image_bbox(xref)
                        bbox = BoundingBox(bbox_rect.x0, bbox_rect.y0, bbox_rect.x1, bbox_rect.y1, page_idx)
                    except Exception:
                        bbox = None

                    section_id = f"page_{page_idx+1}_embed_image_{img_index+1}"
                    elements.append(
                        StructuredElement(
                            section_id=section_id,
                            content_type="image",
                            text="",
                            image_snippets=[img_name],
                            bounding_box=bbox,
                            page=page_idx,
                            notes="Embedded image extracted via PyMuPDF",
                        )
                    )
        finally:
            try:
                doc.close()
            except Exception:
                pass

        return elements

    def _extract_tables_camelot(self, file_path: str) -> List[StructuredElement]:
        """Extract tables using Camelot as a fallback when pdfplumber misses tables."""
        if camelot is None:
            return []

        elements: List[StructuredElement] = []
        try:
            tables = camelot.read_pdf(file_path, pages="all", flavor="stream")
            for idx, table in enumerate(tables):
                try:
                    table_json = table.df.to_json(orient="records")
                except Exception:
                    table_json = table.df.to_json()

                table_text = "\n".join([" | ".join(row) for row in table.df.values.tolist()])
                section_id = f"camelot_table_{idx+1}"
                elements.append(
                    StructuredElement(
                        section_id=section_id,
                        content_type="table",
                        text=table_text,
                        metadata={"table_data": json.loads(table_json)},
                    )
                )
        except Exception:
            pass

        return elements


def parse_uploaded_file(file: BinaryIO, file_name: str, output_dir: str = None) -> List[Dict[str, Any]]:
    """
    Parse an uploaded file into structured content.
    
    Args:
        file: File-like object containing the uploaded file
        file_name: Name of the uploaded file
        output_dir: Directory to store extracted images
        
    Returns:
        List of structured elements
    """
    # Create temp directory if not provided
    if not output_dir:
        output_dir = tempfile.mkdtemp()
        os.makedirs(output_dir, exist_ok=True)
    
    # Reset file pointer in case it has been read earlier
    try:
        file.seek(0)
    except Exception:
        pass

    # Save uploaded file to temp path
    file_ext = Path(file_name).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
        temp_file.write(file.read())
        temp_path = temp_file.name
    
    try:
        # Parse the file
        parser = DocumentParser(output_dir=output_dir)
        elements = parser.parse_file(temp_path)
        
        return elements
    finally:
        # Clean up temp file
        try:
            Path(temp_path).unlink(missing_ok=True)
        except:
            pass 

    # ------------------------------------------------------------------
    # Image preprocessing helper
    # ------------------------------------------------------------------

    def _preprocess_image_for_ocr(self, pil_img: Image.Image) -> Image.Image:
        """Apply grayscale, bilateral blur and adaptive threshold to improve OCR."""
        if cv2 is None:
            return pil_img

        img_cv = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2GRAY)
        # Denoise
        img_cv = cv2.bilateralFilter(img_cv, d=9, sigmaColor=75, sigmaSpace=75)
        # Adaptive threshold
        img_cv = cv2.adaptiveThreshold(img_cv, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                       cv2.THRESH_BINARY, 31, 2)
        return Image.fromarray(img_cv) 